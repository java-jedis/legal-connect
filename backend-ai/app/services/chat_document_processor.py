"""
Chat document processing service for handling documents uploaded during chat sessions
"""

import os
import uuid
import asyncio
from typing import List, Dict, Any, Optional
import logging
from sqlalchemy.orm import Session
import PyPDF2
import docx
from io import BytesIO

from app.utils.text_processing import TextProcessor
from app.services.embeddings import EmbeddingService
from app.services.vectordb import QdrantService
from app.db.models import ChatDocument, ChatDocumentChunk, ChatSession
from app.db.database import get_db

logger = logging.getLogger(__name__)

class ChatDocumentProcessor:
    """Service for processing documents uploaded during chat sessions"""
    
    def __init__(
        self,
        embedding_service: EmbeddingService,
        vectordb_service: QdrantService
    ):
        self.embedding_service = embedding_service
        self.vectordb_service = vectordb_service
        self.text_processor = TextProcessor()
        
        logger.info("Initialized ChatDocumentProcessor")
    
    async def upload_and_process_document(
        self,
        file_content: bytes,
        filename: str,
        file_type: str,
        session_id: str,
        user_id: str,
        db: Session
    ) -> Dict[str, Any]:
        """
        Upload and process a document for a chat session
        
        Args:
            file_content: Raw file content
            filename: Original filename
            file_type: File type (pdf, txt, docx)
            session_id: Chat session ID
            user_id: User ID from JWT
            db: Database session
        
        Returns:
            Processing result
        """
        try:
            # Verify session exists and belongs to user
            session = db.query(ChatSession).filter(
                ChatSession.id == session_id,
                ChatSession.user_id == user_id
            ).first()
            
            if not session:
                raise ValueError(f"Session {session_id} not found or not owned by user {user_id}")
            
            # Generate unique filename (for database record only)
            file_id = str(uuid.uuid4())
            file_extension = os.path.splitext(filename)[1]  # Use os.path.splitext instead of Path
            unique_filename = f"{file_id}{file_extension}"
            
            # Create database record (no file_path since not saving to disk)
            chat_document = ChatDocument(
                session_id=session_id,
                user_id=user_id,  # Add user_id to match session user_id
                filename=unique_filename,
                original_filename=filename,
                file_size=len(file_content),
                file_type=file_type,
                file_path=None,  # No file saved to disk
                processing_status='processing'
            )
            
            db.add(chat_document)
            db.commit()
            db.refresh(chat_document)
            
            # Process document asynchronously
            processing_result = await self._process_document_content(
                chat_document,
                file_content,
                file_type,
                db
            )
            
            return {
                "success": True,
                "document_id": chat_document.id,
                "filename": filename,
                "processing_result": processing_result
            }
            
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _process_document_content(
        self,
        chat_document: ChatDocument,
        file_content: bytes,
        file_type: str,
        db: Session
    ) -> Dict[str, Any]:
        """Process document content and create embeddings"""
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text_content = await self._extract_pdf_text(file_content)
            elif file_type.lower() == 'txt':
                text_content = file_content.decode('utf-8')
            elif file_type.lower() in ['doc', 'docx']:
                text_content = await self._extract_docx_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if not text_content or len(text_content.strip()) < 50:
                raise ValueError("Document contains insufficient text content")
            
            # Update document with extracted text
            chat_document.text_content = text_content
            db.commit()
            
            # Create text chunks
            chunks = self.text_processor.chunk_text(text_content)
            
            if not chunks:
                raise ValueError("No text chunks could be created from document")
            
            # Generate embeddings for chunks
            chunk_texts = [chunk["text"] for chunk in chunks]
            embeddings = await self.embedding_service.generate_embeddings_batch(chunk_texts)
            
            # Prepare documents for vector database with chat-specific namespace
            vector_documents = []
            chunk_records = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                # Create document chunk record
                chunk_record = ChatDocumentChunk(
                    chat_document_id=chat_document.id,
                    chunk_text=chunk["text"],
                    chunk_index=i,
                    page_number=chunk.get("page_number")
                )
                chunk_records.append(chunk_record)
                
                # Prepare for vector database with session-specific metadata
                vector_doc = {
                    "document_id": chat_document.id,
                    "chunk_id": None,  # Will be updated after database commit
                    "text": chunk["text"],
                    "session_id": chat_document.session_id,
                    "document_name": chat_document.original_filename,
                    "document_type": "chat_upload",
                    "metadata": {
                        "chunk_index": i,
                        "length": len(chunk["text"]),
                        "session_id": chat_document.session_id,
                        "file_type": chat_document.file_type,
                        "upload_timestamp": chat_document.created_at.isoformat()
                    }
                }
                vector_documents.append(vector_doc)
            
            # Add chunk records to database
            db.add_all(chunk_records)
            db.commit()
            
            # Refresh chunk records to get their IDs
            for chunk_record in chunk_records:
                db.refresh(chunk_record)
            
            # Update vector document IDs
            for vector_doc, chunk_record in zip(vector_documents, chunk_records):
                vector_doc["chunk_id"] = chunk_record.id
            
            # Add to vector database
            vector_ids = await self.vectordb_service.add_documents(vector_documents, embeddings)
            
            # Update chunk records with vector IDs
            for chunk_record, vector_id in zip(chunk_records, vector_ids):
                chunk_record.vector_id = vector_id
            
            # Update document status
            chat_document.total_chunks = len(chunks)
            chat_document.processing_status = 'completed'
            
            db.commit()
            
            logger.info(f"Successfully processed chat document {chat_document.id} with {len(chunks)} chunks")
            
            return {
                "chunks_created": len(chunks),
                "embeddings_created": len(embeddings),
                "text_length": len(text_content)
            }
            
        except Exception as e:
            logger.error(f"Error processing document content: {e}")
            
            # Update document status to failed
            chat_document.processing_status = 'failed'
            chat_document.processing_error = str(e)
            db.commit()
            
            raise e
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"Page {page_num + 1}:\n{page_text}\n")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def get_session_documents(self, session_id: str, user_id: str, db: Session) -> List[Dict[str, Any]]:
        """Get all documents for a chat session"""
        try:
            # Verify session ownership
            session = db.query(ChatSession).filter(
                ChatSession.id == session_id,
                ChatSession.user_id == user_id
            ).first()
            
            if not session:
                raise ValueError(f"Session {session_id} not found or not owned by user {user_id}")
            
            documents = db.query(ChatDocument).filter(
                ChatDocument.session_id == session_id
            ).all()
            
            result = []
            for doc in documents:
                result.append({
                    "id": doc.id,
                    "filename": doc.original_filename,
                    "file_type": doc.file_type,
                    "file_size": doc.file_size,
                    "total_chunks": doc.total_chunks,
                    "processing_status": doc.processing_status,
                    "processing_error": doc.processing_error,
                    "created_at": doc.created_at.isoformat()
                })
            
            return result
            
        except Exception as e:
            logger.error(f"Error getting session documents: {e}")
            raise e
    
    async def delete_session_document(self, document_id: str, session_id: str, user_id: str, db: Session) -> bool:
        """Delete a document from a chat session"""
        try:
            # Verify session ownership and document exists
            session = db.query(ChatSession).filter(
                ChatSession.id == session_id,
                ChatSession.user_id == user_id
            ).first()
            
            if not session:
                raise ValueError(f"Session {session_id} not found or not owned by user {user_id}")
            
            document = db.query(ChatDocument).filter(
                ChatDocument.id == document_id,
                ChatDocument.session_id == session_id
            ).first()
            
            if not document:
                return False
            
            # Delete vector embeddings
            chunks = db.query(ChatDocumentChunk).filter(
                ChatDocumentChunk.chat_document_id == document_id
            ).all()
            
            vector_ids = [chunk.vector_id for chunk in chunks if chunk.vector_id]
            if vector_ids:
                await self.vectordb_service.delete_documents(vector_ids)
            
            # Delete chunks
            db.query(ChatDocumentChunk).filter(
                ChatDocumentChunk.chat_document_id == document_id
            ).delete()
            
            # Delete document record
            db.delete(document)
            db.commit()
            
            logger.info(f"Deleted chat document {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            db.rollback()
            raise e
    
    async def search_session_documents(
        self,
        query: str,
        session_id: str,
        user_id: str,
        top_k: int = 5,
        db: Session = None
    ) -> List[Dict[str, Any]]:
        """Search within documents of a specific chat session"""
        try:
            # Verify session ownership
            session = db.query(ChatSession).filter(
                ChatSession.id == session_id,
                ChatSession.user_id == user_id
            ).first()
            
            if not session:
                raise ValueError(f"Session {session_id} not found or not owned by user {user_id}")
            
            # Generate query embedding
            query_embedding = await self.embedding_service.generate_query_embedding(query)
            
            # Search with session filter
            results = await self.vectordb_service.search_similar(
                query_embedding=query_embedding,
                top_k=top_k,
                score_threshold=0.7,
                filter_conditions={"session_id": session_id}
            )
            
            return results
            
        except Exception as e:
            logger.error(f"Error searching session documents: {e}")
            raise e
